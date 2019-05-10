# -*- coding: utf-8 -*-
"""
Created on Sat Apr 20 15:21:26 2019

@author: Mehmet
"""
import pyttsx3
import webbrowser
import smtplib
import random
import speech_recognition as sr
import wikipedia
import datetime
import wolframalpha
import time
import os
import sys
from gtts import gTTS 
from translate import Translator
from tkinter import *
from  docx import Document

pencere = Tk()
pencere.title("ASİSTAN")
pencere.geometry("1000x600+425+200")
pencere.state("normal")
pencere.wm_attributes("-alpha", 0.9)

mkullanıcıadı=""
mşifre=""
malıcı=""
mailm=""

engine = pyttsx3.init('sapi5')
en_voice_id = "HKEY_LOCAL_MACHINE\SOFTWARE\Microsoft\Speech\Voices\Tokens\TTS_MS_EN-US_ZIRA_11.0"
engine.setProperty('voice', en_voice_id)
engine.say('')
voices = engine.getProperty('voices')
engine.setProperty('voice', voices[len(voices)-1].id)

client = wolframalpha.Client('7GVLY3-E89JYA3TQ7')

def speaktr(audio):
    tts = gTTS(text=audio, lang='tr')
    tts.save("asd.mp3")
    os.system("asd.mp3")
    
def speak(audio):    
    print('Bilgisayar: ' + audio)
    engine.say(audio)
    engine.runAndWait()

def dinleme():
    r = sr.Recognizer()
    with sr.Microphone() as source:
        print("birsey söyle")
        r.adjust_for_ambient_noise(source)
        girdi = r.listen(source)
        try:
            print("Söylediniz  " + r.recognize_google(girdi,language='TR'))
            son=r.recognize_google(girdi,language='TR')
        except sr.UnknownValueError:
            speak('Üzgünüm efendim! Ben yapamadım,komutu yazmayı deneyin!')
            son = str(input('Komut:'))
        except sr.RequestError as e:
            print("Could not request results from Google Speech Recognition service; {0}".format(e))
    return son
    
def greetMe():
    saat = int(datetime.datetime.now().hour)
    if saat >= 5 and saat < 10:
        speak('Günaydın!')
        
    if saat >= 10 and saat < 12:
        speak('İyi Günler! ')
        
    if saat >= 12 and saat < 13:
        speak('İyi Öğlenler! ')
        
    if saat >= 13 and saat < 17:
        speak('İyi Günler! ')

    if saat >= 17 and saat < 22:
        speak('İyi Akşamlar! ')
        
    if saat >= 22 and saat < 0:
        speak('İyi Geceler! ')

    if saat >= 0 and saat <=5:
        speak('İyi Geceler!')

def açgoogle():
    webbrowser.open('www.google.co.in')

def açyoutube():
    webbrowser.open('www.youtube.com')
    
def açgmail():
    webbrowser.open('www.gmail.com')
    
def havadurumu():
    webbrowser.open('https://www.google.com.tr/search?q=elazığ+hava+durumu&rlz=')
    
def aragoogle(girdi):
    webbrowser.open('www.google.com/search?q='+girdi)
    
def arayoutube(girdi):
    webbrowser.open('www.youtube.com/results?search_query='+girdi)
    
    
def arama(girdi):
    
    query = turing(girdi)
    try:
        try:
            res = client.query(query)
            results = next(res.results).text
            results=ingtur(results)
            metin.delete(0,END)
            metin.insert(0,results)
            speak(results)

        except:
            results = wikipedia.summary(query, sentences=2)
            results=ingtur(results)
            metin.delete(0,END)
            metin.insert(0,results)
            speak(results)

    except:
        aragoogle(girdi)
    

def ingtur(girdi):
    translator= Translator(to_lang="tr")
    translation = translator.translate(girdi)
    return translation

def turing(girdi):
    translator= Translator(from_lang="tr",to_lang="en")
    translation = translator.translate(girdi)
    return translation

def mail():
    
    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.ehlo()
        server.starttls()
        server.login(mkullanıcıadı, mşifre)
        server.sendmail(mkullanıcıadı, malıcı, mailm)
        server.close()
        speak('Email gönderildi!')
    except:
        speak('Üzgünüm efendim! Şu anda mesajınızı gönderemiyorum!')
        
       
def kapan():
    sys.exit()       
        
    
    #D YE KAYDEDİYO
def masaustuword():
    wordkayıtbelgesi=metin.get()
    print(wordkayıtbelgesi)
    document = Document()
    document.add_heading("Asistan Metin Belgesi",level=2)
    document.add_paragraph(wordkayıtbelgesi)
    document.save('D://deneme.docx')


def verialara():
    arama(ara.get())

def verialg():
    aragoogle(arag.get())

def verialy():
    arayoutube(aray.get())
    
def verialm():
    mkullanıcıadı=kullanıcı.get()
    mşifre=şifre.get()
    malıcı=alıcı.get()
    mailm=metin.get()
    print(mkullanıcıadı+mşifre+malıcı+mailm)
    mail()

def müzik():
    speak('Tamam, işte müziğin! eğlenceli')
    filepath = ['D:\\a.mp3', 'D:\\1.mp3' , 'D:\\1.mp3' , 'D:\\a.mp3']
    randomsecim=random.choice(filepath)
    os.startfile(randomsecim)


def dinle():
    h=dinleme().lower()
    
    if 'hava nasıl' in h:
        speak('tamam')     
        havadurumuB.invoke()

    elif 'aç youtube' in h:
        speak('tamam')
        youtubeac.invoke()

    elif 'aç google' in h:
        speak('tamam')
        googleac.invoke()

    elif 'aç gmail' in h:
        speak('tamam')
        mailaç.invoke()
        
    elif 'ara youtube' in h:
        speak('tamam')
        print(h)
        aray.insert(1,"aasd")
        aray.delete(0,END)
        aray.insert(0,h)
        araby.invoke()


    elif 'ara google' in h:
        speak('tamam')
        print(h)
        arag.insert(1,"aasd")
        arag.delete(0,END)
        arag.insert(0,h)
        arabg.invoke()
        
        
    elif "naber" in h or 'nasılsın' in h:
            stMsgs = [' Sadece işimi yapıyorum! ' , ' Ben iyiyim! ' , ' Güzel! ' , ' Ben iyiyim ve enerji doluyum ']
            speak(random.choice(stMsgs))

    elif 'mail gönder' in h:
        speak('tamam')
        mailat.invoke()

    elif ' hiçbir şey' in h or 'iptal' in h or 'dur' in h:
            speak('tamam')
            speak('Hoşçakalın efendim, iyi günler.')
            sys.exit()
           
    elif 'merhaba' in h:
            speak('Merhaba efendim')

    elif 'hoşçakal' in h:
            speak('Hoşçakalın efendim, iyi günler.')
            sys.exit()
                  
    elif 'aç müzik' in h:
        müzikaç.invoke()
        
    else:
       # speak('aranıyor')
        print(h)
        ara.delete(0,END)
        ara.insert(0,h)
        arab.invoke()

                
def metinyaz():
    yazıses=dinleme();
    metinl.delete(0,END)
    metinl.insert(0,yazıses)
    
def metinyazind():
    metinyazısı=metin.get()
    metinlyazısı=metinl.get()
    metinyazısı=metinyazısı+" "+metinlyazısı
    metin.delete(0,END)
    metin.insert(0,metinyazısı)
    
        
metinl = Entry()
metinl.place(relx=0.0,rely=0.1,relheight=0.05,relwidth=0.5)

metinlb= Button(text="Alta kaydet",command=metinyazind)
metinlb.place(relx=0.5,rely=0.1,relheight=0.05,relwidth=0.15)
    

ara=Entry()
ara.place(relx=0.005,rely=0.015,relheight=0.05,relwidth=0.2)

arab= Button(text="Araştırma",command=verialara)
arab.place(relx=0.215,rely=0.0,relheight=0.08,relwidth=0.1)


arag=Entry()
arag.place(relx=0.345,rely=0.015,relheight=0.05,relwidth=0.2)

arabg= Button(text="Google da ara",command=verialg)
arabg.place(relx=0.555,rely=0.0,relheight=0.08,relwidth=0.1)


aray=Entry()
aray.place(relx=0.685,rely=0.015,relheight=0.05,relwidth=0.2)

araby= Button(text="Youtube da ara",command=verialy)
araby.place(relx=0.895,rely=0.0,relheight=0.08,relwidth=0.1)


metin=Entry()
metin.place(relx=0.0,rely=0.15,relheight=0.7,relwidth=0.6)


ekullanıcı = Label(text= "Kullanıcı Adı=",anchor="sw")
ekullanıcı.place(relx=0.7,rely=0.1,relheight=0.05,relwidth=0.15) 

kullanıcı=Entry()
kullanıcı.place(relx=0.7,rely=0.15,relheight=0.05,relwidth=0.2)
 
 
eşifre = Label(text= "Şifre=",anchor="sw")
eşifre.place(relx=0.7,rely=0.22,relheight=0.05,relwidth=0.15)

şifre=Entry(show="*")
şifre.place(relx=0.7,rely=0.27,relheight=0.05,relwidth=0.2)


ealıcı = Label(text= "Alıcının Kullanıcı Adı=",anchor="sw")
ealıcı.place(relx=0.7,rely=0.34,relheight=0.05,relwidth=0.15)

alıcı=Entry()
alıcı.place(relx=0.7,rely=0.39,relheight=0.05,relwidth=0.2)


mailat =Button(text="Mail Gönder", command=verialm)
mailat.place(relx=0.75,rely=0.45,relheight=0.1,relwidth=0.2)


dinleb =Button(text="--Dinleme--", command=dinle)
dinleb.place(relx=0.61,rely=0.56,relheight=0.18,relwidth=0.38)


metinyazb =Button(text="Metin Yaz", command=metinyaz)
metinyazb.place(relx=0.0,rely=0.85,relheight=0.1,relwidth=0.2)


kaydet =Button(text="Kaydet", command=masaustuword)
kaydet.place(relx=0.4,rely=0.85,relheight=0.1,relwidth=0.2)


googleac =Button(text="Google Aç", command=açgoogle)
googleac.place(relx=0.61,rely=0.85,relheight=0.05,relwidth=0.38)


youtubeac =Button(text="Youtube Aç", command=açyoutube)
youtubeac.place(relx=0.61,rely=0.90,relheight=0.05,relwidth=0.38)


havadurumuB =Button(text="Hava Durumu", command=havadurumu)
havadurumuB.place(relx=0.61,rely=0.95,relheight=0.05,relwidth=0.38)


mailaç =Button(text="Mail Aç", command=açgmail)
mailaç.place(relx=0.61,rely=0.8,relheight=0.05,relwidth=0.38)


müzikaç =Button(text="Müzik Aç", command=müzik)
müzikaç.place(relx=0.61,rely=0.75,relheight=0.05,relwidth=0.38)


greetMe()
#speak('Merhaba efendim, ben dijital asistanınız SİM !')
#speak('Size nasıl yardımcı olabilirim? ')
pencere.mainloop()















